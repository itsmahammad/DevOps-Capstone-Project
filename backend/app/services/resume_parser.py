"""
Resume Parser Service - Extracts text and structured data from PDF/DOCX
"""
import re
from pypdf import PdfReader
from docx import Document
from typing import Dict, List, Any, Optional
from app.models.schemas import CandidateInfo, Project, Experience, ExperienceSummary, Education
from app.services.ocr_service import ocr_service


class ResumeParser:
    """Parse resumes and extract structured information"""
    
    # Parsing method constants
    PARSING_STANDARD = "standard"
    PARSING_OCR = "ocr"
    PARSING_OCR_UNAVAILABLE = "ocr_unavailable"
    
    # Regex patterns
    EMAIL_PATTERN = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    PHONE_PATTERN = r'(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}|\+\d{1,3}[-.\s]?\d{6,14}'
    LINKEDIN_PATTERN = r'(?:linkedin\.com/in/|linkedin:?\s*)([a-zA-Z0-9-]+)'
    GITHUB_PATTERN = r'(?:github\.com/|github:?\s*)([a-zA-Z0-9-]+)'
    URL_PATTERN = r'https?://[^\s<>"{}|\\^`\[\]]+'
    
    # Section headers
    SECTION_HEADERS = {
        'experience': ['experience', 'work experience', 'employment', 'work history', 'professional experience', 'career history'],
        'education': ['education', 'academic', 'qualification', 'academics', 'educational background'],
        'skills': ['skills', 'technical skills', 'competencies', 'technologies', 'tech stack', 'expertise'],
        'projects': ['projects', 'personal projects', 'academic projects', 'key projects', 'portfolio'],
        'certifications': ['certifications', 'certificates', 'credentials', 'licenses'],
        'summary': ['summary', 'profile', 'objective', 'about', 'professional summary', 'career objective']
    }
    
    # Action verbs for experience analysis
    ACTION_VERBS = [
        'achieved', 'administered', 'analyzed', 'architected', 'automated',
        'built', 'collaborated', 'configured', 'created', 'delivered',
        'designed', 'developed', 'drove', 'enhanced', 'established',
        'executed', 'implemented', 'improved', 'increased', 'integrated',
        'launched', 'led', 'managed', 'mentored', 'migrated',
        'optimized', 'orchestrated', 'oversaw', 'pioneered', 'planned',
        'reduced', 'refactored', 'resolved', 'scaled', 'secured',
        'spearheaded', 'streamlined', 'supervised', 'transformed', 'upgraded'
    ]
    
    def parse(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        """Main parsing method with OCR fallback for scanned PDFs"""
        # Initialize parsing metadata
        parsing_method = self.PARSING_STANDARD
        ocr_confidence = None
        
        # Extract raw text
        if file_ext == '.pdf':
            raw_text = self._extract_pdf_text(file_path)
            has_tables = self._check_pdf_tables(file_path)
            has_images = self._check_pdf_images(file_path)
            
            # Check if we need OCR fallback (only for PDFs)
            raw_text, parsing_method, ocr_confidence = self._apply_ocr_if_needed(
                file_path, raw_text
            )
        else:
            # DOCX files are always text-based, never OCR
            raw_text = self._extract_docx_text(file_path)
            has_tables = self._check_docx_tables(file_path)
            has_images = self._check_docx_images(file_path)
        
        # Parse sections
        sections = self._identify_sections(raw_text)
        
        # Extract structured data
        candidate = self._extract_candidate_info(raw_text)
        experience = self._extract_experience(raw_text, sections.get('experience', ''))
        projects = self._extract_projects(raw_text, sections.get('projects', ''))
        education = self._extract_education(raw_text, sections.get('education', ''))
        
        return {
            "raw_text": raw_text,
            "candidate": candidate,
            "experience": experience,
            "projects": projects,
            "education": education,
            "sections": sections,
            "formatting": {
                "has_tables": has_tables,
                "has_images": has_images,
                "word_count": len(raw_text.split()),
                "line_count": len(raw_text.split('\n'))
            },
            "parsing_method": parsing_method,
            "ocr_confidence": ocr_confidence
        }
    
    def _apply_ocr_if_needed(
        self, 
        file_path: str, 
        standard_text: str
    ) -> tuple:
        """
        Apply OCR fallback if standard extraction is insufficient
        
        Decision Logic:
        - Text length < 800 characters → OCR
        - Word count < 150 → OCR
        - No email found → OCR
        - No phone found → OCR
        
        Never merges OCR + standard text.
        
        Args:
            file_path: Path to PDF file
            standard_text: Text extracted via pypdf
            
        Returns:
            Tuple of (text, parsing_method, ocr_confidence)
        """
        # Check if OCR service is available
        if not ocr_service.is_available():
            return standard_text, self.PARSING_STANDARD, None
        
        # Quick check for email and phone in standard text
        email_match = re.search(self.EMAIL_PATTERN, standard_text)
        phone_match = re.search(self.PHONE_PATTERN, standard_text)
        
        # Determine if OCR is needed
        if not ocr_service.needs_ocr(
            standard_text, 
            email=email_match.group() if email_match else None,
            phone=phone_match.group() if phone_match else None
        ):
            # Standard extraction is good enough
            return standard_text, self.PARSING_STANDARD, None
        
        # Check if PDF is too large for OCR
        if ocr_service.should_skip_ocr(file_path):
            # PDF has too many pages, skip OCR
            return standard_text, self.PARSING_OCR_UNAVAILABLE, None
        
        # Attempt OCR extraction
        ocr_text, parsing_method, confidence = ocr_service.extract_text_with_ocr(
            file_path
        )
        
        if ocr_text and parsing_method == self.PARSING_OCR:
            # OCR succeeded - replace standard text entirely
            return ocr_text, parsing_method, confidence
        else:
            # OCR failed or unavailable - fall back to standard
            return standard_text, parsing_method, confidence
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using pypdf"""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
        return text
    
    def _check_pdf_tables(self, file_path: str) -> bool:
        """Check if PDF contains tables (potential ATS issue)"""
        # pypdf doesn't have built-in table detection
        # We'll use a heuristic: check for table-like patterns in text
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text() or ""
                # Look for table-like patterns (multiple tabs or consistent spacing)
                lines = text.split('\n')
                table_like_lines = 0
                for line in lines:
                    if line.count('\t') >= 2 or line.count('|') >= 2:
                        table_like_lines += 1
                if table_like_lines > 3:
                    return True
        except:
            pass
        return False
    
    def _check_pdf_images(self, file_path: str) -> bool:
        """Check if PDF contains images"""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                if '/XObject' in page.get('/Resources', {}):
                    xobject = page['/Resources']['/XObject']
                    if xobject:
                        for obj in xobject:
                            if xobject[obj]['/Subtype'] == '/Image':
                                return True
        except:
            pass
        return False
    
    def _check_docx_tables(self, file_path: str) -> bool:
        """Check if DOCX contains tables"""
        try:
            doc = Document(file_path)
            return len(doc.tables) > 0
        except:
            return False
    
    def _check_docx_images(self, file_path: str) -> bool:
        """Check if DOCX contains images"""
        try:
            doc = Document(file_path)
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    return True
        except:
            pass
        return False
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify and extract resume sections"""
        sections = {}
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            section_found = None
            
            for section_type, headers in self.SECTION_HEADERS.items():
                for header in headers:
                    if line_lower == header or line_lower.startswith(header + ':') or line_lower.startswith(header + ' '):
                        section_found = section_type
                        break
                if section_found:
                    break
            
            if section_found:
                if current_section:
                    sections[current_section] = '\n'.join(current_content)
                current_section = section_found
                current_content = []
            elif current_section:
                current_content.append(line)
        
        if current_section:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _extract_candidate_info(self, text: str) -> CandidateInfo:
        """Extract candidate contact information"""
        # Get first few lines for name detection
        lines = text.split('\n')[:10]
        name = None
        
        # Name is usually in the first few lines, all caps or title case
        for line in lines:
            line = line.strip()
            if len(line) > 2 and len(line) < 50:
                # Skip lines that look like addresses, emails, or phone numbers
                if '@' in line or re.search(self.PHONE_PATTERN, line):
                    continue
                if any(word in line.lower() for word in ['resume', 'cv', 'curriculum']):
                    continue
                # Check if it looks like a name (2-4 words, alphabetic)
                words = line.split()
                if 1 <= len(words) <= 4 and all(w.replace('.', '').replace('-', '').isalpha() for w in words):
                    name = line
                    break
        
        # Extract email
        email_match = re.search(self.EMAIL_PATTERN, text)
        email = email_match.group() if email_match else None
        
        # Extract phone
        phone_match = re.search(self.PHONE_PATTERN, text)
        phone = phone_match.group() if phone_match else None
        
        # Extract LinkedIn
        linkedin_match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        linkedin = f"linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else None
        
        # Extract GitHub
        github_match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        github = f"github.com/{github_match.group(1)}" if github_match else None
        
        # Extract location (common patterns)
        location = self._extract_location(text)
        
        return CandidateInfo(
            name=name,
            email=email,
            phone=phone,
            location=location,
            linkedin=linkedin,
            github=github
        )
    
    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location from resume"""
        # Common location patterns
        location_patterns = [
            r'(?:Location|Address|Based in|City)[:\s]+([A-Za-z\s,]+)',
            r'([A-Za-z]+,\s*[A-Z]{2})\s*\d{5}',  # City, ST ZIP
            r'([A-Za-z]+,\s*[A-Za-z\s]+,\s*[A-Za-z]+)',  # City, State, Country
        ]
        
        for pattern in location_patterns:
            match = re.search(pattern, text[:500], re.IGNORECASE)
            if match:
                location = match.group(1).strip()
                if len(location) > 3 and len(location) < 50:
                    return location
        return None
    
    def _extract_experience(self, full_text: str, experience_section: str) -> ExperienceSummary:
        """Extract work experience details"""
        positions = []
        text_to_analyze = experience_section if experience_section else full_text
        
        # Split into potential job entries
        entries = self._split_experience_entries(text_to_analyze)
        
        total_months = 0
        
        for entry in entries[:5]:  # Limit to 5 positions
            exp = self._parse_experience_entry(entry)
            if exp.company or exp.role:
                positions.append(exp)
                # Estimate duration
                if exp.duration:
                    months = self._estimate_duration_months(exp.duration)
                    total_months += months
        
        # Calculate overall quality
        overall_quality = 0
        if positions:
            quality_sum = sum(p.bullet_quality for p in positions)
            overall_quality = quality_sum // len(positions)
        
        return ExperienceSummary(
            total_years=round(total_months / 12, 1),
            total_months=total_months,
            positions=positions,
            overall_quality=overall_quality
        )
    
    def _split_experience_entries(self, text: str) -> List[str]:
        """Split experience section into individual entries"""
        # Common patterns that indicate new entry
        date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}'
        
        lines = text.split('\n')
        entries = []
        current_entry = []
        
        for i, line in enumerate(lines):
            # Check if this line might start a new entry
            has_date = re.search(date_pattern, line, re.IGNORECASE)
            is_company_like = (
                line.strip() and
                len(line.strip()) < 60 and
                (line.strip()[0].isupper() if line.strip() else False)
            )
            
            if has_date and current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = []
            
            current_entry.append(line)
        
        if current_entry:
            entries.append('\n'.join(current_entry))
        
        return entries if entries else [text]
    
    def _parse_experience_entry(self, entry: str) -> Experience:
        """Parse a single experience entry"""
        lines = entry.strip().split('\n')
        
        # Try to identify company and role from first few lines
        company = None
        role = None
        duration = None
        
        for line in lines[:3]:
            line = line.strip()
            if not line:
                continue
            
            # Check for duration pattern
            date_match = re.search(
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–—to]+\s*(?:(Present|Current)|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})',
                line, re.IGNORECASE
            )
            if date_match:
                duration = line
                continue
            
            # Role patterns
            role_keywords = ['engineer', 'developer', 'manager', 'analyst', 'designer', 'lead', 'director', 'intern', 'associate', 'specialist', 'consultant']
            if any(kw in line.lower() for kw in role_keywords) and not role:
                role = line
            elif not company and len(line) > 2:
                company = line
        
        # Analyze bullet points
        bullets = [l for l in lines if l.strip().startswith(('•', '-', '*', '●', '○')) or (l.strip() and l.strip()[0].isdigit() and '.' in l[:3])]
        
        # Count action verbs
        action_count = 0
        has_metrics = False
        for bullet in bullets:
            bullet_lower = bullet.lower()
            for verb in self.ACTION_VERBS:
                if verb in bullet_lower:
                    action_count += 1
                    break
            # Check for metrics
            if re.search(r'\d+%|\$\d+|increased|decreased|reduced|improved by', bullet_lower):
                has_metrics = True
        
        # Calculate bullet quality
        bullet_quality = 0
        if bullets:
            action_ratio = action_count / len(bullets)
            bullet_quality = int(action_ratio * 70)
            if has_metrics:
                bullet_quality += 30
            bullet_quality = min(100, bullet_quality)
        
        return Experience(
            company=company,
            role=role,
            duration=duration,
            description='\n'.join(bullets[:5]) if bullets else None,
            bullet_quality=bullet_quality,
            has_metrics=has_metrics,
            action_verbs_count=action_count
        )
    
    def _estimate_duration_months(self, duration_str: str) -> int:
        """Estimate duration in months from string"""
        if 'present' in duration_str.lower() or 'current' in duration_str.lower():
            # Assume 1 year for current positions
            return 12
        
        # Try to find two dates
        date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*(\d{4})'
        dates = re.findall(date_pattern, duration_str, re.IGNORECASE)
        
        if len(dates) >= 2:
            try:
                start_year = int(dates[0])
                end_year = int(dates[1])
                return max(1, (end_year - start_year) * 12)
            except:
                pass
        
        return 12  # Default to 1 year
    
    def _extract_projects(self, full_text: str, projects_section: str) -> List[Project]:
        """Extract project information"""
        projects = []
        text_to_analyze = projects_section if projects_section else full_text
        
        # Split into project entries
        entries = self._split_project_entries(text_to_analyze)
        
        for entry in entries[:5]:  # Limit to 5 projects
            project = self._parse_project_entry(entry)
            if project.title:
                projects.append(project)
        
        return projects
    
    def _split_project_entries(self, text: str) -> List[str]:
        """Split projects section into individual entries"""
        lines = text.split('\n')
        entries = []
        current_entry = []
        
        for line in lines:
            # New project typically starts with a title-like line
            if line.strip() and not line.strip().startswith(('•', '-', '*', '●')):
                if current_entry and len(current_entry) > 1:
                    entries.append('\n'.join(current_entry))
                    current_entry = []
            current_entry.append(line)
        
        if current_entry:
            entries.append('\n'.join(current_entry))
        
        return entries
    
    def _parse_project_entry(self, entry: str) -> Project:
        """Parse a single project entry"""
        lines = entry.strip().split('\n')
        
        title = None
        technologies = []
        description = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            if i == 0 or not title:
                # First non-empty line is likely title
                title = line.replace('•', '').replace('-', '').strip()
            else:
                # Check for tech keywords
                tech_match = re.search(r'(?:Tech|Technologies|Built with|Stack)[:\s]+(.+)', line, re.IGNORECASE)
                if tech_match:
                    techs = tech_match.group(1).split(',')
                    technologies.extend([t.strip() for t in techs])
                else:
                    description.append(line)
        
        # Calculate project score
        score = 50  # Base score
        if technologies:
            score += 20
        if description:
            score += 15
            # Check for impact keywords
            desc_text = ' '.join(description).lower()
            if any(w in desc_text for w in ['improved', 'increased', 'reduced', 'users', 'revenue']):
                score += 15
        
        return Project(
            title=title,
            technologies=technologies,
            description='\n'.join(description) if description else None,
            impact=None,
            score=min(100, score)
        )
    
    def _extract_education(self, full_text: str, education_section: str) -> List[Education]:
        """Extract education information"""
        education = []
        text_to_analyze = education_section if education_section else full_text
        
        lines = text_to_analyze.split('\n')
        current_edu = {}
        
        degree_keywords = ['bachelor', 'master', 'phd', 'doctorate', 'b.s.', 'b.a.', 'm.s.', 'm.a.', 'mba', 'b.tech', 'm.tech', 'b.e.', 'm.e.', 'diploma', 'associate']
        
        for line in lines:
            line_lower = line.lower().strip()
            if not line_lower:
                continue
            
            # Check for degree
            for keyword in degree_keywords:
                if keyword in line_lower:
                    if current_edu:
                        education.append(Education(**current_edu))
                    current_edu = {'degree': line.strip()}
                    break
            else:
                # Check for year
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                if year_match and current_edu:
                    current_edu['year'] = year_match.group()
                
                # Check for GPA
                gpa_match = re.search(r'(?:GPA|CGPA)[:\s]*(\d+\.?\d*)', line, re.IGNORECASE)
                if gpa_match and current_edu:
                    current_edu['gpa'] = gpa_match.group(1)
                
                # Institution might be next line after degree
                if current_edu and 'institution' not in current_edu and len(line.strip()) > 3:
                    if not any(kw in line_lower for kw in degree_keywords):
                        current_edu['institution'] = line.strip()
        
        if current_edu:
            education.append(Education(**current_edu))
        
        return education[:3]  # Limit to 3 entries
